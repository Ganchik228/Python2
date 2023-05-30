import tkinter as tk
from tkinter import messagebox
from rectangle import Rectangle
from triangle import Triangle
from trapezoid import Trapezoid
import docx


class CalculatorGUI:
    def __init__(self, master):
        self.master = master
        master.title("Calculator")

        self.label_shape = tk.Label(master, text="Choose figure:")
        self.label_shape.grid(row=0, column=0)

        self.shape_var = tk.StringVar(master)
        self.shape_var.set("Rectangle")

        self.shape_menu = tk.OptionMenu(
            master, self.shape_var, "Rectangle", "Triangle", "Trapezoid", command=self.show_inputs)
        self.shape_menu.grid(row=0, column=1)

        self.label_input1 = tk.Label(master)
        self.label_input1.grid(row=1, column=0)

        self.entry_input1 = tk.Entry(master)
        self.entry_input1.grid(row=1, column=1)

        self.label_input2 = tk.Label(master)
        self.label_input2.grid(row=2, column=0)

        self.entry_input2 = tk.Entry(master)
        self.entry_input2.grid(row=2, column=1)

        self.label_input3 = tk.Label(master)
        self.label_input3.grid(row=3, column=0)

        self.entry_input3 = tk.Entry(master)
        self.entry_input3.grid(row=3, column=1)

        self.button_calculate = tk.Button(
            master, text="Calculate", command=self.calculate)
        self.button_calculate.grid(row=4, column=1)

        self.button_save = tk.Button(
            master, text="Save results", command=self.save_results)
        self.button_save.grid(row=4, column=0)

        self.result_label = tk.Label(master)
        self.result_label.grid(row=5, column=0, columnspan=2)

    def show_inputs(self, shape):
        if shape == "Rectangle":
            self.label_input1.config(text="Length:")
            self.label_input2.config(text="Width:")
            self.label_input3.config(text="")
            self.entry_input3.delete(0, tk.END)
            self.entry_input3.config(state="disabled")
        elif shape == "Trianlge":
            self.label_input1.config(text="Side A:")
            self.label_input2.config(text="Side B:")
            self.label_input3.config(text="Side C:")
            self.entry_input3.config(state="normal")
        elif shape == "Trapezoid":
            self.label_input1.config(text="Base A:")
            self.label_input2.config(text="Base B:")
            self.label_input3.config(text="Height:")
            self.entry_input3.config(state="normal")

    def calculate(self):
        shape = self.shape_var.get()

        if shape == "Rectangle":
            try:
                length = float(self.entry_input1.get())
                width = float(self.entry_input2.get())
                rectangle = Rectangle(length, width)
                area = rectangle.area()
                perimeter = rectangle.perimeter()
                diagonal = rectangle.diagonal()
                inscribed_radius = rectangle.inscribed_radius()
                circumscribed_radius = rectangle.circumscribed_radius()
                result_text = f"S: {area:.2f}, P: {perimeter:.2f}, Diagonal: {diagonal:.2f}, Inscribed circle R: {inscribed_radius:.2f}, Circumscribed circle R: {circumscribed_radius:.2f}"
                self.result_label.config(text=result_text)
            except ValueError:
                messagebox.showerror("Error", "Input numbers")

        elif shape == "Triangle":
            try:
                a = float(self.entry_input1.get())
                b = float(self.entry_input2.get())
                c = float(self.entry_input3.get())
                triangle = Triangle(a, b, c)
                area = triangle.area()
                perimeter = triangle.perimeter()
                circumradius = triangle.circumradius()
                inradius = triangle.inradius()
                result_text = f"S: {area:.2f}, P: {perimeter:.2f}, Inscribed circle R: {inradius:.2f}, Circumscribed circle R: {circumradius:.2f}"
                self.result_label.config(text=result_text)
            except ValueError:
                messagebox.showerror("Error", "Input numbers")

        elif shape == "Trapezoid":
            try:
                a = float(self.entry_input1.get())
                b = float(self.entry_input2.get())
                h = float(self.entry_input3.get())
                trapezoid = Trapezoid(a, b, h)
                area = trapezoid.area()
                perimeter = trapezoid.perimeter()
                inscribed_radius = trapezoid.inscribed_radius()
                circumscribed_radius = trapezoid.circumscribed_radius()
                result_text = f"S: {area:.2f}, P: {perimeter:.2f}, Inscribed circle R: {inscribed_radius:.2f}, Circumscribed circle R: {circumscribed_radius:.2f}"
                self.result_label.config(text=result_text)
            except ValueError:
                messagebox.showerror("Error", "Input numbers")

    def save_results(self):
        shape = self.shape_var.get()
        result_text = self.result_label.cget("text")
        if result_text:
            doc = docx.Document()
            doc.add_heading(f"{shape} Results", 0)
            doc.add_paragraph(result_text)
            doc.save(f"{shape} Results.docx")
            messagebox.showinfo("Inforamtion", "Results saved in file")


root = tk.Tk()
app = CalculatorGUI(root)
root.mainloop()